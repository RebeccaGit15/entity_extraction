from gliner import GLiNER
from docx import Document
from docx.shared import RGBColor

model = GLiNER.from_pretrained("urchade/gliner_multi-v2.1")
model_2 = GLiNER.from_pretrained("urchade/gliner_multi")
model_3 = GLiNER.from_pretrained("urchade/gliner_multi_pii-v1")

text = """
استنادا للأمر الشفهي لمدير عمليات قيادة الجيش و نفيدكم أنه بتاريخه الساعه 18:15 وبالتنسيق مع مديرية المخابرات أرسلت دورية معززة من ل.م.12 الى بلدة كوشا/عكار وعملت على مداهمة عدد من المنازل لاشخاص مطلوبين عل الشكل التالي :
1- المطلوب خالد الزين والدته حنان مواليد 1991 رقم السجل 84 كوشا حيث تم توقيفه كونه مطلوب بموجب وثيقة إطلاق نار ومطلوب بعدة وثائق أعمال ارهابية واطلاق نار وقذائف ار بي جي اخرها الوثيقة رقم 6 تاريخ 14/4/2014.
2- المطلوب ----- والدته ----- مواليد ---- رقم السجل ----.
3- المطلوب ----- والدته ----- مواليد ---- رقم السجل ----.
4- المطلوب ----- والدته ----- مواليد ---- رقم السجل ----.
5- المواطن ------ والدته ----- مواليد ---- رقم السجل ----.
تم العثور على مضبوطات مذكور بالائحة المرفقة ربطا بالوثيقة. كما تمت مداهمة مبنى عائدة للمطلوب خالد الزين (المذكور أعلاه) يستأجره عدد من السوريين حيث اوقف داخل المبنى كل من :
1- السوري ------ والدته ------ مواليد ----.
2- السوري ------ والدته ------ مواليد ----.
3- السوري ------ والدته ------ مواليد ----.
لدخولهم الاراضي البنانية خلسة.
سلم الموقوفين مع أغراضهم الخاصة والمضبوطات الى الشرطة العسكرية باستثناء رمانة يدوية سلمت الى سرية الهندسة ليصار اجراء الازم بشأنها لاحقا.
"""

color_map = {
    "Person": RGBColor(164, 92, 204),
    "Mother": RGBColor(255, 0, 0), 
    "Parent": RGBColor(255, 0, 0), 
    "Grenade": RGBColor(0, 153, 76),
    "Guns": RGBColor(0, 0, 102),
    "Weapon": RGBColor(0, 0, 102),
    "Missile": RGBColor(0, 153, 76),
    "Family of": RGBColor(35, 40, 249),
    "Family name": RGBColor(35, 40, 249),
    "Organization": RGBColor(35, 40, 249),
    "Family_member": RGBColor(35, 40, 249),
}

questions_and_labels = [
    ("Who are the persons mentioned?", ["Person"]),
    ("Who are the mothers mentioned?", ["Mother"]),
    ("Who are the parents mentioned?", ["Parent"]),
    ("What are the grenades mentioned?", ["Grenade"]),
    ("What are the guns mentioned?", ["Guns"]),
    ("What are the weapons mentioned?", ["Weapon"]),
    ("What are the missiles mentioned?", ["Missile"]),
    ("What are the family mentioned?", ["Family of"]),
    ("What are the family names mentioned?", ["Family name"]),
    ("What are the organizations mentioned?", ["Organization"]),
    ("What are the family members mentioned?", ["Family_member"])
]

def split_text_with_overlap(text, max_length=500, overlap=50):
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), max_length - overlap):
        chunk = " ".join(words[i:i + max_length])
        chunks.append(chunk)
    
    return chunks

def ask_and_extract_with_chunks(model, text, questions_and_labels, max_length=500, overlap=50):
    chunks = split_text_with_overlap(text, max_length, overlap)
    all_entities = []
    seen_entities = set()
    
    for idx, chunk in enumerate(chunks):
        print(f"\nProcessing chunk {idx+1}/{len(chunks)}...\n")
        for question, labels in questions_and_labels:
            print(f"Question: {question}")
            entities = model.predict_entities(chunk, labels, threshold=0.5)
            for entity in entities:
                entity_text = entity['text']
                entity_position = text.find(entity_text)
                
                if entity_position != -1 and (entity_position, entity_position + len(entity_text)) not in seen_entities:
                    print(f"Entity: {entity['text']}, Label: {entity['label']}, Score: {entity['score']}")
                    all_entities.append(entity)
                    
                    seen_entities.add((entity_position, entity_position + len(entity_text)))
            print("\n")
    
    return all_entities

def save_entities_to_docx(text, entities, filename="output.docx"):
    doc = Document()
    
    doc.add_heading('The extracted entities are:', level=2)
    
    for entity in entities:
        entity_text = entity['text']
        entity_label = entity['label']
        entity_score = entity['score']
        color = color_map.get(entity_label, RGBColor(0, 0, 0)) 
        
        para = doc.add_paragraph()
        run = para.add_run(f"Entity: {entity_text}, Label: {entity_label}, Score: {entity_score:.2f}")
        run.font.color.rgb = color
    
    doc.add_paragraph() 

    para = doc.add_paragraph()
    
    last_position = 0
    
    entities_sorted = sorted(entities, key=lambda e: text.find(e['text']))
    
    for entity in entities_sorted:
        entity_text = entity['text']
        entity_label = entity['label']
        color = color_map.get(entity_label, RGBColor(0, 0, 0))
        
        entity_position = text.find(entity_text, last_position)
        
        para.add_run(text[last_position:entity_position])
        
        run = para.add_run(entity_text)
        run.font.color.rgb = color
        
        last_position = entity_position + len(entity_text)
    
    para.add_run(text[last_position:])
    
    doc.save(filename)

entities = ask_and_extract_with_chunks(model, text, questions_and_labels)
entities_2 = ask_and_extract_with_chunks(model_2, text, questions_and_labels)
entities_3 = ask_and_extract_with_chunks(model_3, text, questions_and_labels)

# Display the extracted entities in the terminal
# display_entities_in_terminal(entities)

save_entities_to_docx(text, entities, "results_gliner_multi-v2.1/0.5/توقيفات.docx")
save_entities_to_docx(text, entities_2, "results_gliner_multi/0.5/توقيفات.docx")
save_entities_to_docx(text, entities_3, "results_gliner_multi_pii-v1/0.5/توقيفات.docx")
